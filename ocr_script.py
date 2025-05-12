import os
from docx import Document
from io import BytesIO
from PIL import Image
import pytesseract
import re

# Specify the paths
doc_path = r"C:\Users\DESKTOP.GE\Desktop\pyphoto\ChangePy.docx"  # Path to your Word document
folder_path = r"C:\Users\DESKTOP.GE\Desktop\pyphoto\photos_folder"  # Folder to save images
output_path = r"C:\Users\DESKTOP.GE\Desktop\pyphoto\output.docx"  # Path for saving the output Word document

# Ensure the folder exists, create it if it doesn't
os.makedirs(folder_path, exist_ok=True)

# Function to extract and save images from a Word document
def extract_images_to_folder(doc_path, folder_path):
    doc = Document(doc_path)
    image_count = 0  # To keep track of image numbers

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:  # Check if the part is an image
            image_count += 1
            image_data = rel.target_part.blob  # Get image data
            image = Image.open(BytesIO(image_data))  # Open image from bytes
            
            # Define the image path and save it
            image_path = os.path.join(folder_path, f"image_{image_count}.png")
            image.save(image_path, "PNG")  # Save as PNG
            
            print(f"Saved {image_path}")

    print(f"Total images extracted: {image_count}")

# Function to perform OCR on an image and return extracted text
def image_to_text(image_path):
    # Open the image using PIL
    image = Image.open(image_path)
    # Use pytesseract to extract text from the image
    text = pytesseract.image_to_string(image, config='--psm 6')  # PSM 6 for a single uniform block of text
    return text

# Function to process images in a folder and extract text tables
def process_images_in_folder(folder_path):
    tables = []
    for filename in os.listdir(folder_path):
        if filename.endswith(".png"):  # Check for image file types
            image_path = os.path.join(folder_path, filename)
            text = image_to_text(image_path)  # Call the OCR function
            print(f"Extracted text from {filename}:\n{text}\n")  # Display the extracted text
            
            # Append the extracted text to the tables list
            tables.append(text.strip())
    
    return tables

# Function to save extracted text into a Word document with dynamic tables
def save_tables_to_word(tables, output_path):
    doc = Document()
    
    for i, text in enumerate(tables):
        # Split text into lines for table construction
        lines = text.strip().split('\n')
        
        if lines:
            # Create a new table with dynamic columns based on the first line
            first_line_columns = re.split(r'\s{2,}', lines[0])  # Splitting by two or more spaces to determine columns
            num_columns = len(first_line_columns)

            # Create the table with the number of columns found
            table = doc.add_table(rows=1, cols=num_columns)
            table.style = 'Table Grid'
            
            # Add headers from the first line
            hdr_cells = table.rows[0].cells
            for j, header in enumerate(first_line_columns):
                hdr_cells[j].text = header.strip()

            # Fill the table with the remaining lines
            for line in lines[1:]:
                if line.strip():  # Skip empty lines
                    # Using regex to separate each entry based on variable spaces
                    columns = re.split(r'\s{2,}', line)  # Split by 2 or more spaces
                    row_cells = table.add_row().cells
                    for j, col in enumerate(columns):
                        if j < num_columns:  # Prevent index error
                            row_cells[j].text = col.strip()
        
        doc.add_paragraph(f"\nExtracted from image {i + 1}\n")  # Add a paragraph to separate tables
    
    doc.save(output_path)
    print(f"Tables saved to {output_path}")

# Run the functions
extract_images_to_folder(doc_path, folder_path)
tables = process_images_in_folder(folder_path)
save_tables_to_word(tables, output_path)
